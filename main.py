from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles

from pydantic import BaseModel
import shutil
import os
from pathlib import Path
from docx import Document
from datetime import datetime

app = FastAPI()

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 정적 파일 마운트
app.mount("/static", StaticFiles(directory="static"), name="static")

# 업로드된 파일 정보 저장을 위한 리스트
uploaded_files_info = []

# 루트 페이지
@app.get("/")
async def root():
    return FileResponse('static/index.html')


# '내 자소서 쓰기' 페이지
@app.get("/write")
async def write_page():
    return FileResponse('static/write.html')

@app.on_event("startup")
async def load_existing_files():
    upload_dir = "uploaded_files"
    
    # 폴더가 존재하면 파일 목록을 가져와서 업데이트
    if os.path.exists(upload_dir):
        uploaded_files_info.clear()  # 기존 목록을 초기화
        for filename in os.listdir(upload_dir):
            file_path = os.path.join(upload_dir, filename)
            if os.path.isfile(file_path):
                # 파일이 존재하면 목록에 추가
                upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                uploaded_files_info.append({
                    "filename": filename,
                    "upload_time": upload_time
                })

@app.post("/delete_file")
async def delete_file(filename: str = Form(...)):
    upload_dir = "uploaded_files"
    file_path = os.path.join(upload_dir, filename)

    # 파일이 존재하는지 확인
    if not os.path.exists(file_path):
        return {"error": "File not found"}

    try:
        # 파일 삭제
        os.remove(file_path)
    except Exception as e:
        return {"error": f"Failed to delete file: {str(e)}"}

    # uploaded_files_info 리스트에서도 파일 정보 삭제
    uploaded_files_info[:] = [file for file in uploaded_files_info if file["filename"] != filename]

    return {"message": "File deleted successfully"}

# 워드 파일 업로드 및 처리
@app.post("/process_word")
async def process_word_file(file: UploadFile = File(...)):
    # 업로드된 파일을 저장할 폴더 경로
    upload_dir = "uploaded_files"
    
    # 디렉터리가 없으면 생성
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir)

    # 파일 확장자 확인
    if not file.filename.endswith((".docx", ".doc")):
        return {"error": "Invalid file type. Please upload a .docx or .doc file."}

    # 파일 이름 중복 방지 - 같은 이름이 있으면 (1), (2) 등을 붙임
    file_location = f"{upload_dir}/{file.filename}"
    base, ext = os.path.splitext(file_location)
    counter = 1

    while os.path.exists(file_location):
        file_location = f"{base} ({counter}){ext}"
        counter += 1

    # 업로드한 파일을 임시로 저장
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # 파일 이름과 업로드 날짜 기록
    upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    uploaded_files_info.append({
        "filename": os.path.basename(file_location),  # 파일 이름에 번호가 붙은 경우를 처리
        "upload_time": upload_time
    })

    # 워드 파일 내용 읽기
    doc = Document(file_location)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # 워드 파일 내용을 반환
    return {"filename": os.path.basename(file_location), "content": "\n".join(full_text)}

# 업로드된 파일 목록 제공 API
@app.get("/uploaded_files")
async def get_uploaded_files():
    return {"files": uploaded_files_info}



# 파일 이름 변경 API
@app.post("/rename_file")
async def rename_file(old_filename: str = Form(...), new_filename: str = Form(...)):
    upload_dir = "uploaded_files"
    old_file_path = os.path.join(upload_dir, old_filename)
    new_file_path = os.path.join(upload_dir, new_filename)

    # 파일이 존재하는지 확인
    if not os.path.exists(old_file_path):
        return {"error": "File not found"}

    # 새로운 파일 이름이 이미 존재할 경우, 숫자를 추가하여 고유하게 만듦
    base, ext = os.path.splitext(new_file_path)
    counter = 1
    while os.path.exists(new_file_path):
        new_file_path = f"{base} ({counter}){ext}"
        counter += 1

    # 파일 이름 변경
    os.rename(old_file_path, new_file_path)

    # 리스트에서도 파일 이름 변경
    for file_info in uploaded_files_info:
        if file_info["filename"] == old_filename:
            file_info["filename"] = os.path.basename(new_file_path)

    return {"message": "File renamed successfully", "new_filename": os.path.basename(new_file_path)}

# 파일 내용 저장 API
@app.post("/save_file_content")
async def save_file_content(filename: str = Form(...), content: str = Form(...)):
    file_location = f"uploaded_files/{filename}"
    
    # 파일이 존재하는지 확인
    if not os.path.exists(file_location):
        return {"error": "File not found"}

    # 워드 파일 내용을 업데이트
    doc = Document()
    # 하나의 패러그래프로 텍스트를 처리하고 줄바꿈(\n)을 유지
    paragraph = doc.add_paragraph()
    for line in content.split("\n"):
        paragraph.add_run(line)
        # paragraph.add_run("\n")  # 줄바꿈을 그대로 추가

    doc.save(file_location)

    return {"message": "File content saved successfully"}


# 업로드된 파일 내용 반환 API
@app.get("/file_content/{filename}")
async def get_file_content(filename: str):
    # 파일 경로 설정
    file_location = f"uploaded_files/{filename}"
    
    # 파일이 존재하는지 확인
    if not os.path.exists(file_location):
        return {"error": "File not found"}

    # 워드 파일 내용 읽기
    doc = Document(file_location)
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 파일 내용을 반환
    return {"filename": filename, "content": "\n".join(full_text)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
